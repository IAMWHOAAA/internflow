from io import BytesIO

import pytest
from django.core.files.uploadedfile import SimpleUploadedFile
from docx import Document

from jobs.services import resume_analysis
from jobs.services.resume_analysis import (
    ResumeExtractionError,
    analyze_resume,
    extract_keywords,
    extract_resume_text,
)


def make_docx(text: str, name: str = "resume.docx") -> SimpleUploadedFile:
    document = Document()
    for line in text.splitlines():
        document.add_paragraph(line)
    output = BytesIO()
    document.save(output)
    return SimpleUploadedFile(
        name,
        output.getvalue(),
        content_type=("application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    )


def test_extracts_text_from_docx_paragraphs_and_tables():
    document = Document()
    document.add_paragraph("教育经历：江苏科技大学 软件工程")
    document.add_paragraph("项目经历：InternFlow 求职工作台")
    document.add_paragraph("负责使用 Django 设计并实现投递管理、权限隔离和自动化测试。")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Python"
    table.cell(0, 1).text = "Django"
    output = BytesIO()
    document.save(output)
    uploaded = SimpleUploadedFile("resume.docx", output.getvalue())

    text = extract_resume_text(uploaded)

    assert "江苏科技大学" in text
    assert "Python | Django" in text


def test_extracts_text_from_pdf_pages(monkeypatch):
    class Page:
        def extract_text(self):
            return "项目经历 " + "使用 Python Django 完成开源项目并优化性能 " * 5

    class Reader:
        is_encrypted = False
        pages = [Page()]

        def __init__(self, uploaded_file):
            self.uploaded_file = uploaded_file

    monkeypatch.setattr(resume_analysis, "PdfReader", Reader)
    uploaded = SimpleUploadedFile("resume.pdf", b"fake pdf", content_type="application/pdf")

    assert "开源项目" in extract_resume_text(uploaded)


def test_rejects_encrypted_and_textless_pdf(monkeypatch):
    class EncryptedReader:
        is_encrypted = True
        pages = []

        def __init__(self, uploaded_file):
            pass

        def decrypt(self, password):
            raise RuntimeError

    monkeypatch.setattr(resume_analysis, "PdfReader", EncryptedReader)
    encrypted = SimpleUploadedFile("resume.pdf", b"encrypted")

    with pytest.raises(ResumeExtractionError, match="密码"):
        extract_resume_text(encrypted)

    short_docx = make_docx("太短")
    with pytest.raises(ResumeExtractionError, match="文字太少"):
        extract_resume_text(short_docx)


def test_rejects_unknown_or_broken_file():
    unknown = SimpleUploadedFile("resume.txt", b"plain text")
    with pytest.raises(ResumeExtractionError, match="PDF"):
        extract_resume_text(unknown)

    broken = SimpleUploadedFile("resume.docx", b"not a zip")
    with pytest.raises(ResumeExtractionError, match="无法读取"):
        extract_resume_text(broken)


def test_analyzes_rich_resume_with_explainable_signals():
    text = """
    教育经历
    江苏科技大学 软件工程
    专业技能
    Python Django SQL Docker Git Linux
    项目经历
    InternFlow 项目：设计并实现 Django 求职工作台，支持 20 项测试。
    负责数据库索引优化，使请求耗时降低 35%，服务 100 个用户。
    实习经历
    开发数据分析平台，使用 Pandas 完成 5 项报表，效率提升 50%。
    """

    result = analyze_resume(text)

    assert result["score"] >= 80
    assert result["sections"] == {
        "education": True,
        "projects": True,
        "experience": True,
        "skills": True,
    }
    assert {"Python", "Django", "SQL", "Docker", "Git", "Linux", "数据分析"} <= set(
        result["keywords"]
    )
    assert result["metrics_count"] >= 4
    assert any("量化" in strength for strength in result["strengths"])


def test_analyzes_sparse_and_overlong_resume():
    sparse = analyze_resume("Python " + "学习课程和完成作业 " * 20)

    assert sparse["score"] < 70
    assert any("项目经历" in weakness for weakness in sparse["weaknesses"])
    assert any("量化" in weakness for weakness in sparse["weaknesses"])

    long_text = "教育经历 项目经历 实习经历 专业技能 Python " + "开发优化 " * 1000
    overlong = analyze_resume(long_text)

    assert any("较长" in weakness for weakness in overlong["weaknesses"])


def test_extract_keywords_is_case_insensitive():
    assert extract_keywords("FASTAPI + PostgreSQL + React + HTTP") == [
        "FastAPI",
        "React",
        "SQL",
        "HTTP",
    ]
