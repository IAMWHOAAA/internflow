import re
from pathlib import Path

from docx import Document
from pypdf import PdfReader


class ResumeExtractionError(ValueError):
    """Raised when a resume cannot be safely converted to useful text."""


KEYWORD_ALIASES = {
    "Python": ("python",),
    "Django": ("django",),
    "Flask": ("flask",),
    "FastAPI": ("fastapi",),
    "Java": ("java",),
    "Spring": ("spring", "spring boot"),
    "C++": ("c++",),
    "JavaScript": ("javascript", "js"),
    "TypeScript": ("typescript", "ts"),
    "Vue": ("vue", "vue.js"),
    "React": ("react", "react.js"),
    "SQL": ("sql", "mysql", "postgresql", "sqlite"),
    "Redis": ("redis",),
    "Docker": ("docker", "容器"),
    "Linux": ("linux",),
    "Git": ("git", "github"),
    "机器学习": ("机器学习", "machine learning"),
    "深度学习": ("深度学习", "deep learning"),
    "数据分析": ("数据分析", "data analysis", "pandas"),
    "测试": ("自动化测试", "单元测试", "pytest", "测试开发"),
    "HTTP": ("http", "restful", "rest api"),
}

SECTION_MARKERS = {
    "education": ("教育经历", "教育背景", "学历", "education"),
    "projects": ("项目经历", "项目经验", "个人项目", "projects", "project experience"),
    "experience": ("实习经历", "工作经历", "实践经历", "experience", "employment"),
    "skills": ("专业技能", "技能清单", "技术栈", "skills", "technical skills"),
}

ACTION_WORDS = (
    "负责",
    "实现",
    "设计",
    "开发",
    "优化",
    "搭建",
    "主导",
    "重构",
    "提升",
    "降低",
    "完成",
)


def _normalize_text(text: str) -> str:
    text = text.replace("\x00", " ")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)


def extract_resume_text(uploaded_file) -> str:
    """Extract text from a PDF or DOCX without sending it to a third party."""
    extension = Path(uploaded_file.name).suffix.lower()
    uploaded_file.seek(0)
    try:
        if extension == ".pdf":
            reader = PdfReader(uploaded_file)
            if reader.is_encrypted:
                try:
                    reader.decrypt("")
                except Exception as error:
                    raise ResumeExtractionError("暂不支持带密码的 PDF 简历。") from error
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif extension == ".docx":
            document = Document(uploaded_file)
            parts = [paragraph.text for paragraph in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            text = "\n".join(parts)
        else:
            raise ResumeExtractionError("只支持 PDF 或 DOCX 格式。")
    except ResumeExtractionError:
        raise
    except Exception as error:
        raise ResumeExtractionError("无法读取这份简历，请确认文件没有损坏。") from error
    finally:
        uploaded_file.seek(0)

    text = _normalize_text(text)
    if len(text) < 80:
        raise ResumeExtractionError("提取到的文字太少。扫描版 PDF 请先进行 OCR，或上传 DOCX 版本。")
    return text


def extract_keywords(text: str) -> list[str]:
    lower_text = text.lower()
    detected = []
    for keyword, aliases in KEYWORD_ALIASES.items():
        if any(alias.lower() in lower_text for alias in aliases):
            detected.append(keyword)
    return detected


def analyze_resume(text: str) -> dict:
    """Create an explainable resume review from structure and evidence signals."""
    lower_text = text.lower()
    sections = {
        key: any(marker in lower_text for marker in markers)
        for key, markers in SECTION_MARKERS.items()
    }
    keywords = extract_keywords(text)
    metrics = re.findall(
        r"(?:\d+(?:\.\d+)?%|\d+\s*(?:人|次|项|个|万|ms|秒|天|用户|请求|行))",
        text,
        re.IGNORECASE,
    )
    action_count = sum(text.count(word) for word in ACTION_WORDS)
    token_count = len(re.findall(r"[\u4e00-\u9fff]|[A-Za-z0-9+#.]+", text))

    score = 30
    score += 10 if sections["education"] else 0
    score += 15 if sections["projects"] else 0
    score += 15 if sections["experience"] else 0
    score += 10 if sections["skills"] else 0
    score += min(len(metrics), 4) * 3
    score += 5 if 350 <= token_count <= 1600 else 0
    score += 3 if action_count >= 4 else 0
    score = min(score, 100)

    strengths = []
    weaknesses = []
    suggestions = []

    present_sections = [key for key, present in sections.items() if present]
    if len(present_sections) >= 3:
        strengths.append("简历结构较完整，招聘者能快速找到核心信息。")
    if sections["projects"]:
        strengths.append("包含项目经历，具备展开技术面试的素材。")
    if sections["experience"]:
        strengths.append("包含实践或实习经历，能够证明真实协作经验。")
    if len(metrics) >= 2:
        strengths.append(f"识别到 {len(metrics)} 处量化结果，成果表达比较有说服力。")
    if len(keywords) >= 5:
        strengths.append(f"技术关键词覆盖较丰富，已识别 {len(keywords)} 项能力。")
    if action_count >= 4:
        strengths.append("使用了较多行动动词，个人贡献表达相对清晰。")

    missing_labels = {
        "education": "教育经历",
        "projects": "项目经历",
        "experience": "实践/实习经历",
        "skills": "技能清单",
    }
    for key, label in missing_labels.items():
        if not sections[key]:
            weaknesses.append(f"未清晰识别到“{label}”板块。")
            suggestions.append(f"增加独立的“{label}”标题，并按时间倒序组织。")

    if len(metrics) < 2:
        weaknesses.append("项目和经历中的量化成果偏少。")
        suggestions.append("为关键经历补充性能、规模、效率或结果数据。")
    if action_count < 4:
        weaknesses.append("个人贡献与具体动作不够突出。")
        suggestions.append("使用“设计、实现、优化、负责”等动词说明你亲自完成的工作。")
    if token_count < 350:
        weaknesses.append("简历信息密度偏低，可能不足以支撑深入追问。")
        suggestions.append("补充项目背景、技术决策、个人贡献和最终结果。")
    elif token_count > 1600:
        weaknesses.append("简历内容较长，关键信息可能被稀释。")
        suggestions.append("压缩重复描述，把每段经历保留在 3-5 个高价值要点内。")
    if len(keywords) < 3:
        weaknesses.append("可识别的技术关键词较少。")
        suggestions.append("在技能和项目描述中明确写出实际使用过的技术栈。")

    if not strengths:
        strengths.append("简历已具备可读取的文本基础，可以继续针对目标岗位优化。")
    if not weaknesses:
        weaknesses.append("整体信号较完整，下一步应针对具体岗位调整关键词顺序。")
        suggestions.append("把与目标岗位最相关的技能和项目移动到更靠前的位置。")

    if score >= 85:
        summary = "结构和证据都比较完整，重点转向岗位定制与面试表达。"
    elif score >= 70:
        summary = "基础质量良好，补强量化结果和个人贡献后会更有竞争力。"
    elif score >= 55:
        summary = "核心内容已经具备，但结构与成果证据仍有明显提升空间。"
    else:
        summary = "目前更像信息草稿，建议先补齐关键板块再进行岗位定制。"

    return {
        "score": score,
        "summary": summary,
        "strengths": strengths[:5],
        "weaknesses": weaknesses[:5],
        "suggestions": suggestions[:6],
        "sections": sections,
        "keywords": keywords,
        "metrics_count": len(metrics),
        "word_count": token_count,
    }
